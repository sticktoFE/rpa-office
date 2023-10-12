from functools import partial
import os
from pathlib import Path
import re
from PySide6.QtCore import (
    Qt,
    QEvent,
    QTimer,
    QStandardPaths,
    Slot,
    QRect,
    QPoint,
    QRegularExpression,
)
from PySide6.QtGui import (
    QShortcut,
    QKeySequence,
    QTextDocument,
    QTextCharFormat,
    QColor,
    QTextCursor,
    QFont,
)
from PySide6.QtWidgets import (
    QApplication,
    QDialog,
    QFileDialog,
    QTableView,
    QHeaderView,
    QListWidgetItem,
    QAbstractItemView,
    QAbstractScrollArea,
    QSlider,
)
from mytools.video.OCRVideo import OCRVideo
from myutils.GeneralQThread import Worker
from myutils.image_convert import cv2pixmap
from myutils.info_out_manager import get_temp_file, get_temp_folder
from route.OCRRequest import get_ocr_result
from ui.component.TipsShower import TipsShower
from .OCRResult import Ui_OCRResult
import pandas as pd
import numpy as np
import cv2
from myutils.office_tools import to_excel_auto_column_weight
from ui.component.pandas_model import PandasModel
from docx import Document
import sqlparse
from PIL import Image
from natsort import natsorted
from PySide6.QtCore import QThreadPool

# 综合截屏界面
from mytools.screen_shot.MainLayer import MainLayer


class TotalMessage(QDialog, Ui_OCRResult):
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.adjustSize()
        self.setWindowFlags(
            Qt.WindowType.WindowMinMaxButtonsHint | Qt.WindowType.WindowCloseButtonHint
        )  # Qt.WindowStaysOnTopHint|
        # self.setWindowFlag(Qt.FramelessWindowHint)  # 没有窗口栏
        # self.setWindowFlag(Qt.Tool)  # 不然exec_执行退出后整个程序退出
        self.setAttribute(Qt.WA_TranslucentBackground)  # 设置背景透明
        self.select_file_btn.clicked.connect(self.select_files)
        self.start_scan_btn.clicked.connect(self.screen_shot_info)
        self.start_btn_video.clicked.connect(self.start_ocr_video)
        self.start_btn_video_screenshot.clicked.connect(self.ocr_video_screenshot)
        self.start_btn_video_monitor.clicked.connect(self.ocr_video_monitor)
        self.start_btn_video_stop.clicked.connect(self.ocr_video_close)
        self.format_table.clicked.connect(self.display_data_table)
        self.export_data_excel.clicked.connect(self.save_data_excel)
        self.export_data_word.clicked.connect(self.save_data_word)
        self.find_button.clicked.connect(self.find_text)
        self.replace_button.clicked.connect(self.replace_text_cache)
        self.find_shortcut = QShortcut(
            QKeySequence(Qt.Modifier.CTRL + Qt.Key.Key_F), self
        )
        self.find_shortcut.activated.connect(self.show_find_dialog)
        # 设置快捷键 Ctrl+/
        self.comment_shortcut = QShortcut(QKeySequence("Ctrl+/"), self)
        self.comment_shortcut.activated.connect(self.add_separator)
        # 设置快捷键 Ctrl+Shift+F
        self.format_shortcut = QShortcut(
            QKeySequence(Qt.Modifier.CTRL + Qt.Modifier.SHIFT + Qt.Key.Key_F), self
        )
        self.format_shortcut.activated.connect(self.format_sql_text)
        self.file_list_widget.installEventFilter(self)
        self.file_list_widget.setAcceptDrops(True)
        self.file_list_widget.itemClicked.connect(self.show_item_ocr_content)
        self.file_list_widget.setHorizontalScrollBarPolicy(
            Qt.ScrollBarPolicy.ScrollBarAlwaysOff
        )  # Disable horizontal scroll bar
        self.file_list_widget.setVerticalScrollMode(
            QAbstractItemView.ScrollMode.ScrollPerPixel
        )  # Set scroll mode to pixel-based scrolling
        self.file_list_widget.setSizeAdjustPolicy(
            QAbstractScrollArea.SizeAdjustPolicy.AdjustToContents
        )
        self.file_list_widget.adjustSize()
        self.graphicsView.installEventFilter(self)
        self.graphicsView.setMouseTracking(True)
        self.gridEdit.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.ocrTextEdit.installEventFilter(self)
        self.ocrTextEdit.textChanged.connect(self.sync_text_edit)
        self.slider.setMinimum(0)
        self.slider.setMaximum(100)
        self.slider.setTickPosition(QSlider.TickPosition.TicksRight)
        self.slider.setTickInterval(1)
        self.slider.valueChanged.connect(self.re_ocr_current_image)
        # 设置分割的两个窗口初始大小相等
        self.splitter.setSizes([1, 1])
        # 弹窗提示信息显示器
        self.Tipsshower = TipsShower("  ", parent=self, fontsize=20)
        self.screenshot = MainLayer()
        self.screenshot.screen_shot_pic_signal.connect(self.init_info)
        self.screenshot.close_signal.connect(self.show)
        # 存储打开文件获得的OCR结果
        self.file_ocr_result = {}  # Store OCR results for each image
        self.upscaled_img = {}
        # 存储点击截屏获得的OCR结果
        self.screen_shot_ocr_result = None
        self.highlight_rect = None
        self.file_path = None

    def init_info(self, content):
        self.screen_shot_ocr_result = content
        self.show()
        if type(content) is np.ndarray:
            show = cv2.resize(content, (320, 260))
            self.graphicsView.setPixmap(cv2pixmap(show))
            # 以下是使用QTextEdit方式插入图片，这个会不断往下追加，先用QLabel吧
            # document = self.videoEdit.document()
            # cursor = QTextCursor(document)
            # p1 = cursor.position() # returns int
            # cursor.insertImage(showImage)
        elif type(content) is tuple:
            # 识别的图，上面带识别框，即文字区域划线
            self.src_long_pic = cv2pixmap(content[0])
            pic_draw = content[1]
            self.graphicsView.setPixmap(cv2pixmap(pic_draw))
            ocr_result = self.format_ocr_info(content[2])
            self.ocrTextEdit.setText(ocr_result)
            # 如果原生带格式识别，就展示出来
            ocr_result_html = content[3]
            if len(ocr_result_html) > 0:
                self.htmlEdit.show()
                self.htmlEdit.clear()
                for html_script in ocr_result_html:
                    self.htmlEdit.append(html_script)
            else:
                self.htmlEdit.hide()
        elif type(content) is str:
            self.ocrTextEdit.setText(ocr_result)
        # self.cameraLabel.resize(480,260)

    # 通过快捷键 ctrl+/增加---------
    def add_separator(self):
        cursor = self.ocrTextEdit.textCursor()
        selected_text = cursor.selectedText()
        lines = selected_text.split("\u2029")  # 前台返回的换行符不是\n
        new_text = "\u2029".join(["---------" + line for line in lines])
        cursor.insertText(new_text)

    def screen_shot_info(self):
        self.hide()
        # 使用 QTimer 实现延迟操作
        QTimer.singleShot(500, self.screenshot.screen_shot)

    @Slot()
    def on_re_screen_shot_clicked(self):
        # 使用 QTimer 实现延迟操作
        QTimer.singleShot(500, self.screenshot.screen_shot)

    def display_data_table(self):
        self.view = QTableView()
        self.view.resize(1200, 800)
        self.view.resizeColumnsToContents()
        self.view.horizontalHeader().setStretchLastSection(True)
        self.view.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)
        self.view.setAlternatingRowColors(True)
        self.view.setSelectionBehavior(QTableView.SelectRows)
        # self.view.resizeRowsToContents()
        # QTableWidget.resizeRowsToContents(self.view)
        model = PandasModel(self.content_list)
        self.view.setModel(model)
        self.view.show()
        # self.hide()

    def save_data_excel(self):
        temp_path = get_temp_folder(execute_file_path=__file__)
        file_path = get_temp_file(des_folder=temp_path)
        i = 0
        for content_sec in self.content_list:
            if type(content_sec) is pd.DataFrame:
                to_excel_auto_column_weight(
                    {f"sheet{str(i)}": content_sec}, file_name=file_path
                )
                i += 1
        # 打开文件
        if i != 0:
            os.startfile(file_path)
        # 选择文件进行pdf内容识别

    def save_data_word(self):
        file, _ = QFileDialog.getSaveFileName(
            self, "保存到", "example.docx", "word files(*.docx *doc)"
        )
        if file:
            # 创建 Document 对象，等价于在电脑上打开一个 Word 文档
            document = Document()
            # 在 Word 文档中添加一个标题
            document.add_heading("这是一个标题", level=0)
            # 文档添加段落
            p = document.add_paragraph("这是白给的段落")
            # 添加带样式的文字
            # 添加段落，文本可以包含制表符（\t）、换行符（\n）或回车符（\r）等
            # add_run() 在段落后面追加文本
            p.add_run("\n我倾斜了").italic = True  # 添加一个倾斜文字
            # 添加一个加粗文字
            p.add_run(f"\n{self.gridEdit.toPlainText()}").bold = True
            # 保存文档
            document.save(file)
        # self.close()

    # 视频扫描
    def start_ocr_video(self):
        if not hasattr(self, "video_ocr"):
            self.video_ocr = OCRVideo()
            self.video_ocr.signal.connect(self.ocr_video_result)
            self.video_ocr.signal_img.connect(self.ocr_video_result)
            # # 准备就绪，启动线程开始扫描 此时 run中
            self.video_ocr.start()
        if not self.video_ocr.isRunning():
            self.video_ocr.start()

    def ocr_video_screenshot(self):
        self.video_ocr.pause()

    def ocr_video_monitor(self):
        self.video_ocr.resume()

    def ocr_video_close(self):
        self.video_ocr.stop()

    def ocr_video_result(self, result):
        # print(result)
        self.init_info(result)

    # 保存图片到硬盘
    def save_local(self):
        file, _ = QFileDialog.getSaveFileName(
            self, "保存到", "screenshot.jpg", "Image files(*.jpg *.gif *.png)"
        )
        if file:
            shot_img = self.graphicsView.pixmap.toImage()
            shot_img.save(file)
        self.close()

    ## 保存图片到硬盘---没有上面的简洁
    # 作用：复制图片
    # def copyPic(self):
    #     with open(self.graphicsView.pixmap, "rb") as f_src:
    #         # 读取图片内容并存储到content变量
    #         content = f_src.read()
    #         with open("./1.jpg", "wb") as f_copy:
    #             # 源图片的内容以二进制形式写入新图片
    #             f_copy.write(content)

    # 保存图片到内存
    def save_to_clipboard(self):
        clipboard = QApplication.clipboard()
        clipboard.setPixmap(self.graphicsView.pixmap)
        self.close()

    def closeEvent(self, event):
        if hasattr(self, "video_ocr") and self.video_ocr:
            self.video_ocr.stop()

    def open_image(self):
        """
        select image file and open it
        :return:
        """
        img_name, _ = QFileDialog.getOpenFileName(
            self, "Open Image File", "*.jpg;;*.png;;*.jpeg"
        )
        self.box.set_image(img_name)

    # from transformers import pipeline

    # 中文信息脱敏
    def process_text(self, input_text):
        # 使用正则表达式识别所有中文字符
        chinese_chars = re.findall(r"[\u4e00-\u9fff]+", input_text)
        # 创建中文字符串到占位符的映射字典
        chinese_to_placeholder = {}
        placeholder = "__CHINESE_TEXT__"
        masked_text = input_text
        for i, chinese_char in enumerate(chinese_chars):
            if chinese_char not in chinese_to_placeholder:
                placeholder_i = f"{placeholder}_{i}"
                masked_text = masked_text.replace(chinese_char, placeholder_i)
                chinese_to_placeholder[chinese_char] = placeholder_i
        # 使用大模型进行文本处理和修正
        # text_generator = pipeline("text-generation", model="your-model-name")
        # corrected_text = text_generator(masked_text)[0]["generated_text"]
        corrected_text = masked_text
        # 将修正后的文本中的占位符恢复为原来的中文
        for chinese_char, placeholder_i in chinese_to_placeholder.items():
            corrected_text = corrected_text.replace(placeholder_i, chinese_char)
        return corrected_text

    def handling_exception(self, text):
        text = re.sub(r"，", ",", text)
        text = re.sub(r"；", ";", text)
        text = re.sub(r"θ", "0", text)
        text = re.sub(r"(\s*?_|_\s*?)", "_", text)
        text = re.sub(r"e\s*?[）)]", "0)", text)
        text = re.sub(r"( >\s*?o|>\s*?e)", ">0", text)
        text = re.sub(r"to\.", "t0.", text)
        text = re.sub(r" _", "_", text)
        text = re.sub(r"[”‘′‘’·\"`．“]", "", text)
        # 把in等后面的中文括号换成英文括号
        text = re.sub(
            r"(in|count|sum|nvl|nv1|and|or|to_char|to_date)\s*?（", r" \1 (", text
        )
        # 把then等后面的中文括号换成英文括号
        text = re.sub(r"）\s*?(then|and|or)", r") \1 ", text)
        # <>= *' sss 替换成 <>='sss'
        text = re.sub(r"([<>=]+)\s*?[*']+?\s*?(\w+)[\s*']*", r"\1'\2' ", text)
        text = re.sub(r"([<>=]+)\s*?[*']*?\s*?(\w+)\s*?[*']+?", r"\1'\2' ", text)
        # - *' sss 替换成 ='sss'
        text = re.sub(r"([-])\s*?[*']+?\s*?(\w+)\s*?[*']*", r"='\2' ", text)
        text = re.sub(r"([-])\s*?(\w+)\s*?[*']+?", r"='\2' ", text)
        text = re.sub(r"\s*?[*]+?(\w+)\s*?[*']+?", r"'\1' ", text)
        text = re.sub(r"[*']\s+?(\w+)\s+?[*']", r"'\1' ", text)
        # 使用正则表达式匹配模式，将匹配的中文产品名称加上单引号
        pattern = r"(?<![\u4e00-\u9fa5-])+[\u4e00-\u9fa5]+[\u4e00-\u9fa5a-zA-Z]*\s*[\(（][^.=]+?[\)）]|(?<![\u4e00-\u9fa5-])+[\u4e00-\u9fa5]+[\u4e00-\u9fa5a-zA-Z\s]*?"
        text = re.sub(pattern, lambda match: f"'{match.group(0)}'", text)
        return text

    def format_sql_text(self):
        text = self.gridEdit.toPlainText()
        text = sqlparse.format(text, reindent=True, keyword_case="lower")
        self.gridEdit.setPlainText(text)

    # 选择图片或（pdf后期支持）内容识别
    def select_files(self):
        default_directory = QStandardPaths.writableLocation(
            QStandardPaths.StandardLocation.PicturesLocation
        )  # 设置默认目录的路径
        files, _ = QFileDialog.getOpenFileNames(
            self, "选择文件", default_directory, "Images (*.jpg *.png);;All Files (*)"
        )
        if files:
            self.process_files(files)

    # 重新识别图片，不断放大图片直到全部识别完图片中的内容
    def re_ocr_current_image(self, value):
        self.re_screen_shot.setEnabled(False)
        image_path = self.file_list_widget.currentItem().data(32)
        if image_path in self.upscaled_img:
            upscaled_img = self.upscaled_img[image_path]
        else:
            upscaled_img = Image.open(image_path)
            self.upscaled_img[image_path] = upscaled_img
        get_ocr_result(
            upscaled_img,
            partial(self.display_ocr_result, image_path),
            adjust_para=value,
            ocr_method="structure",
        )
        # 保存处理后的图像
        # upscaled_img.save("upscaled_xxx.png")

    # 批量处理导入的文件
    def process_files(self, file_paths):
        self.htmlEdit.hide()
        self.file_path = None
        self.file_ocr_result.clear()
        self.file_list_widget.clear()
        # 更新文件列表，下面开始识别并拼接识别的内容
        file_paths = natsorted(file_paths)
        self.files_num = len(file_paths)
        for file_path_str in file_paths:
            file_path = Path(file_path_str)
            # 1、把文件展示到文件列表
            # 创建一个QListWidgetItem，显示文件名，但存储文件路径
            item = QListWidgetItem(file_path.stem)
            item.setData(32, file_path_str)  # 使用UserRole 32 存储文件路径
            self.file_list_widget.addItem(item)
            # 2、开始识别文件
            file_extension = file_path.suffix.lower()
            if file_extension in [".jpg", ".png"]:
                image = Image.open(file_path)
                get_ocr_result(
                    image,
                    partial(self.display_ocr_result, file_path_str),
                    ocr_method="structure",
                )
            elif file_extension == ".txt":
                with open(file_path, "r") as f:
                    self.gridEdit.append(f.read() + "\n\n")

    # 格式化返回的ocr信息列表成字符串
    def format_ocr_info(self, ocr_result: list):
        # 识别的内容列表每个元素是一个tuple,格式为:(left_top_x, left_top_y, right_bottom_x, right_bottom_y, ocr_content)
        str_content_list = [
            react_content_tuple[4] for react_content_tuple in ocr_result
        ]
        return "\n".join(str_content_list)

    # 点击文件列表，展示响应图片和识别内容
    def show_item_ocr_content(self, item):
        # 将Slider的值设置回初始值（这里设置为0）
        # self.slider.setValue(0)
        self.re_screen_shot.setEnabled(True)
        self.file_path = item.data(32)
        ocr_pixmap_draw, ocr_text, _ = self.file_ocr_result[self.file_path]
        # # 阻止textChanged信号
        self.ocrTextEdit.blockSignals(True)
        self.ocrTextEdit.clear()
        # ocr_text = self.handling_exception(ocr_text)
        self.ocrTextEdit.setPlainText(ocr_text)
        # 解除textChanged信号的阻止
        self.ocrTextEdit.blockSignals(False)
        self.graphicsView.setPixmap(ocr_pixmap_draw)
        # 汇总识别结果显示到最终显示框中
        self.find_text(Path(self.file_path).stem)

    # 展示ocr的图片和结果
    def display_ocr_result(self, file_path, result):
        ocr_pixmap_draw = cv2pixmap(result[1])
        # 返回到的识别内容是列表，每一元素代表图片中一行文字
        ocr_react_text_list = result[2]
        ocr_text = self.format_ocr_info(ocr_react_text_list)
        ocr_text = self.handling_exception(ocr_text)
        # 调整图形多增加4次识别，寻找5次识别内容最多者为最优识别,后期再实现
        # for i in range(10, 70, 5):
        #     scale_img = self.scale_image(file_path, self.generate_random(i))
        # Store OCR result
        self.file_ocr_result[file_path] = [
            ocr_pixmap_draw,
            ocr_text,
            ocr_react_text_list,  # 原始ocr信息
        ]
        self.ocrTextEdit.clear()
        self.ocrTextEdit.setPlainText(ocr_text)
        self.graphicsView.setPixmap(ocr_pixmap_draw)
        self.files_num = self.files_num - 1
        # 所有文件识别完成后默认选中第一项
        if self.files_num == 0:
            self.file_list_widget.setCurrentItem(self.file_list_widget.item(0))
            self.show_item_ocr_content(self.file_list_widget.item(0))
            # 汇总识别结果显示到最终显示框中
            self.final_total_ocr_info(Path(file_path).stem)
        elif self.files_num < 0:  # 调整图片，重新识别的情况
            # 汇总识别结果显示到最终显示框中
            self.final_total_ocr_info(Path(file_path).stem)
        self.re_screen_shot.setEnabled(True)

    # 汇总ocr最终结果并滚动到目前选中文件的相应内容部分
    def final_total_ocr_info(self, purpose_text):
        self.gridEdit.clear()
        sorted_keys = natsorted(self.file_ocr_result.keys())
        final_text = "\n".join(
            [
                f"----->{Path(key).stem}\n{self.file_ocr_result[key][1]}"
                for key in sorted_keys
            ]
        )
        # final_text = self.handling_exception(final_text)
        self.gridEdit.setPlainText(final_text)
        self.find_text(purpose_text)

    # 编辑图片ocr的内容时同步展示汇总框里的内容
    def sync_text_edit(self):
        # 把通过截屏重新ocr或人工编辑的结果存起来
        if self.file_path is None:
            return
        self.file_ocr_result[self.file_path][1] = self.ocrTextEdit.toPlainText()
        # 把通过截屏重新ocr的画框图片和识别原始结果存起来
        if self.screen_shot_ocr_result is not None:
            self.file_ocr_result[self.file_path][0] = cv2pixmap(
                self.screen_shot_ocr_result[1]
            )
            self.file_ocr_result[self.file_path][2] = self.screen_shot_ocr_result[2]
            self.screen_shot_ocr_result = None
        self.final_total_ocr_info(Path(self.file_path).stem)
        # 相应内容保存硬盘（保存较慢，拖累界面，所以开子线程）
        worker = Worker(self.update_text_file)
        worker.communication.finished.connect(partial(self.Tipsshower.setText, "信息已缓存"))
        QThreadPool.globalInstance().start(worker)

    def show_find_dialog(self):
        selected_text = self.ocrTextEdit.textCursor().selectedText()
        self.find_input.setText(selected_text)
        self.replace_input.setText(selected_text)
        # 将焦点放到替换框
        self.replace_input.setFocus()
        self.replace_input.selectAll()
        self.find_text(selected_text)

    # 在textedit编辑器中支持查找替换
    def find_text(self, search_text):
        if not search_text:
            search_text = self.find_input.text()
        cursor = self.gridEdit.textCursor()
        format = QTextCharFormat()
        # format.setBackground(self.gridEdit.palette().highlight())
        format.setForeground(QColor("white"))
        regex = QRegularExpression(QRegularExpression.escape(search_text))
        cursor = self.gridEdit.document().find(
            search_text,
            cursor,
            QTextDocument.FindFlag.FindCaseSensitively,  # QTextDocument.FindFlag.FindWholeWords,
        )
        if not cursor.isNull():
            cursor.mergeCharFormat(format)
            self.gridEdit.setTextCursor(cursor)
            # auto scroll to purpose text for gridEdit
            # Calculate the desired vertical position
            target_block_position = (
                self.gridEdit.document()
                .findBlock(cursor.position())
                .layout()
                .position()
            )
            viewport_height = self.gridEdit.viewport().height()
            desired_position = target_block_position.y() - (viewport_height / 2)
            # Scroll to the desired position
            self.gridEdit.verticalScrollBar().setValue(
                self.gridEdit.verticalScrollBar().value() + desired_position
            )
            self.gridEdit.ensureCursorVisible()
        else:
            # 搜索到最下面了，把光标放到文本开头，从头开始搜索
            cursor.movePosition(cursor.MoveOperation.Start)
            self.gridEdit.setTextCursor(cursor)

    def replace_text(self):
        search_text = self.find_input.text()
        replace_text = self.replace_input.text()
        cursor = self.gridEdit.textCursor()
        cursor.beginEditBlock()
        format = QTextCharFormat()
        # format.setBackground(self.gridEdit.palette().highlight())
        format.setForeground(QColor("white"))
        regex = QRegularExpression(QRegularExpression.escape(search_text))
        # Replace all occurrences of search_text with replace_text
        last_cursor = None
        while not cursor.isNull() and not cursor.atEnd():
            cursor = self.gridEdit.document().find(
                regex, cursor, QTextDocument.FindFlag.FindWholeWords
            )
            if not cursor.isNull():
                cursor.insertText(replace_text)
                cursor.mergeCharFormat(format)
                last_cursor = QTextCursor(cursor)
        cursor.endEditBlock()
        # 移动光标到最后一个替换的位置
        if last_cursor:
            self.gridEdit.setTextCursor(last_cursor)

    # 直接替换缓存中的内容
    def replace_text_cache(self):
        search_text = self.find_input.text()
        replace_text = self.replace_input.text()
        # 1、循环替换缓存中的内容
        # 遍历字典的值，对字符串进行替换
        for key, value in self.file_ocr_result.items():
            value[1] = value[1].replace(search_text, replace_text)
        # 2、汇总到显示框，并突出显示替换的内容，方便查看
        self.final_total_ocr_info(Path(self.file_path).stem)
        # 3、重新显示当前图片的内容
        self.ocrTextEdit.setText(self.file_ocr_result[self.file_path][1])

    # 同步保存到硬盘，免得关闭程序丢失加工的文件
    def update_text_file(self):
        # 当前点击图片的画框图片保存
        self.file_ocr_result[self.file_path][0].toImage().save(
            str(
                Path(self.file_path).parent.joinpath(
                    f"draw_{Path(self.file_path).stem}.png"
                )
            )
        )
        # 当前点击图片的识别内容保存到硬盘
        with open(
            Path(self.file_path).parent.joinpath(f"{Path(self.file_path).stem}.sql"),
            "w",
            encoding="utf-8",
        ) as f:
            f.write(self.file_ocr_result[self.file_path][1])

        # 最终汇总结果保存到硬盘
        with open(
            Path(self.file_path).parent.joinpath("final_result.sql"),
            "w",
            encoding="utf-8",
        ) as f:
            f.write(self.gridEdit.toPlainText())

    def eventFilter(self, source, event):
        # 双击拷贝图片
        if (
            event.type() == QEvent.Type.MouseButtonDblClick
            and event.button() == Qt.MouseButton.LeftButton
        ):
            QApplication.clipboard().setPixmap(self.graphicsView.pixmap)
        elif (
            event.type() == QEvent.Type.MouseButtonDblClick
            and event.button() == Qt.MouseButton.RightButton
        ):
            QApplication.clipboard().setPixmap(self.src_long_pic)
            # QApplication.clipboard().setImage(QImage(f"{Path(__file__).parent.parent}/dw/DataFactory/screen_shot/tmp/long_screenshot.png"))
        # elif event.type() == QEvent.Enter:
        #     """
        #     鼠标悬停，提示信息
        #     """
        #     self.Tipsshower.setText("双击保存")
        elif source == self.file_list_widget and event.type() == QEvent.Type.DragEnter:
            mime_data = event.mimeData()
            if mime_data.hasUrls():
                file_paths = [url.toLocalFile() for url in mime_data.urls()]
                self.process_files(file_paths)
                return True
        # elif source == self.ocrTextEdit and event.type() == QEvent.Type.FocusOut:
        #     # print(event.type())
        #     # 连接失去焦点事件到保存函数,改成了textChanged事件
        #     self.sync_text_edit()
        elif (
            source == self.graphicsView
            and event.type() == QEvent.Type.MouseButtonPress
            and event.button() == Qt.MouseButton.LeftButton
        ):
            # 点击图片相应位置显示相应识别内容
            # 获取鼠标点击的坐标
            def mapMouseToImage(click_point):
                # Map the mouse click coordinates to image coordinates
                # return click_point - self.graphicsView.pixmap.rect().topLeft()
                image_label_size = self.graphicsView.size()
                pixmap_size = self.graphicsView.pixmap.size()
                scaled_click_x = (
                    click_point.x() * pixmap_size.width() / image_label_size.width()
                )
                scaled_click_y = (
                    click_point.y() * pixmap_size.height() / image_label_size.height()
                )
                return QPoint(scaled_click_x, scaled_click_y)

            click_point = mapMouseToImage(event.pos())
            # 更新要突出显示的区域
            # print(f"x，y鼠标点击位置{click_point}")
            # 检查是否与OCR结果位置重叠，并突出显示相应的OCR结果
            if self.file_path is not None:
                _, _, ocr_react_texts = self.file_ocr_result[self.file_path]
                highlighted_results = []
                for ocr_react_text in ocr_react_texts:
                    ocr_rect = QRect(
                        ocr_react_text[0],
                        ocr_react_text[1],
                        ocr_react_text[2] - ocr_react_text[0],
                        ocr_react_text[3] - ocr_react_text[1],
                    )
                    if ocr_rect.contains(click_point):
                        highlighted_results.append(ocr_react_text[4])
                if highlighted_results:
                    highlighted_text = "\n".join(highlighted_results)
                    # 可以把提示信息放在鼠标附近
                    # relative_pos = QCursor.pos()
                    # 计算浮窗的位置
                    # relative_pos = self.ocrTextEdit.pos()
                    # absolute_pos = self.mapToGlobal(relative_pos)
                    # 也可以把提示信息放在父窗口附近
                    absolute_pos = self.pos()
                    absolute_pos.setY(
                        absolute_pos.y()
                        if absolute_pos.y() - 30 < 0
                        else absolute_pos.y() - 30
                    )  # 紧贴 Label 上边
                    self.Tipsshower.setText(
                        highlighted_text,
                        pos=(absolute_pos.x(), absolute_pos.y()),
                        autoclose=False,
                    )
                    # self.Tipsshower.setStyleSheet("color:blue")
            self.update()  # 触发重绘事件
        # elif event.type() == QEvent.MouseMove and event.button() == Qt.LeftButton:
        #     """
        #     mouse move events for the widget
        #     :param event: QMouseEvent
        #     :return:
        #     """
        #     self.move_dis = event.pos()- self.start_pos
        #     self.point_x = self.point_x + self.move_dis
        #     self.start_pos = event.pos()
        #     self.repaint()
        # elif event.type() == QEvent.MouseButtonRelease and event.button() == Qt.LeftButton:
        #     """
        #     mouse release events for the widget
        #     :param e: QMouseEvent
        #     :return:
        #     """
        #     self.left_click = False
        return super().eventFilter(source, event)

    # def paintEvent(self, e):
    #     """
    #     receive paint events
    #     :param e: QPaintEvent
    #     :return:
    #     """
    #     scaled_img = self.cameraLabel.pixmap()

    #     self.cameraLabel.rect()
    #     if scaled_img:
    #         painter = QPainter()
    #         painter.begin(self)
    #         painter.drawPixmap(QRect(self.point[0],self.point[1],self.cameraLabel.size().width(),self.cameraLabel.size().height()), scaled_img)
    #         painter.end()
