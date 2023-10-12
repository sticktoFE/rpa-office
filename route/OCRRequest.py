from paddleocr import draw_ocr, to_excel

# from paddleocr.ppstructure.recovery.recovery_to_doc import sorted_layout_boxes
# from paddleocr.ppstructure.recovery.table_process import HtmlToDocx
# from paddleocr.ppstructure.table.predict_table import to_excel
import json
import numpy as np
from copy import deepcopy
import cv2
from docx import Document, shared
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from numpy import vstack
from docx2pdf import convert
from myutils.DateAndTime import clock
from multiprocessing import Pipe
from myutils.image_adjust import scale_image
from route.common import format_result_with_pos, get_pos_contents
from model.OCRServer import OCRServer, OCRServerProcess
import logging
import time
from configparser import ConfigParser
import requests
from myutils.image_convert import get_shot_bytes
from functools import partial
from myutils import globalvar, image_split
from myutils.GeneralQThread import Worker
from PySide6.QtCore import QThreadPool
from PIL import Image


# 1、正常线程调用
class OCRRequest:
    def __init__(self, output="../output", **kwargs):
        self.ocr_server = OCRServer()
        self.output = output

    def set_output_filepath(self, output):
        self.output = output

    @clock
    def ocr(self, cv2_img):
        # print(f"OCRRequest in  {QThread.currentThread()}线程中")
        cv2_img_c = cv2_img.copy()
        result = self.ocr_server.ocr(cv2_img_c)
        # 如果不保持图片原始布局，就只要识别的内容
        returnResult = get_pos_contents(result, from_source="ocr", draw_img=cv2_img_c)
        returnResult = [str(line[2]) for line in returnResult]
        # time.sleep(10)  # 制造阻塞
        return cv2_img_c, returnResult
        # 把图片识别出内容,组合成有效信息返回

    @clock
    def ocr_structure(
        self,
        oneimg,
        keep_format="2",
        only_layout=False,
    ):
        return_result = []
        html_content = []
        # oneimg = np.ascontiguousarray(oneimg)
        # draw_img = cv2.UMat(draw_img).get()
        oneimg_c = oneimg.copy()
        if only_layout:
            return_result = self.ocr_server.ocr_structure(
                oneimg_c, only_layout=only_layout
            )
            print(return_result)
            return oneimg_c, return_result, html_content
        result = self.ocr_server.ocr_structure(
            oneimg_c, return_ocr_result_in_table=True
        )
        # 由于识别的结果 可能顺序乱了，要做个排序
        result.sort(key=lambda x: (x["bbox"][1], x["bbox"][0]))  # 先按y方向排序 在x方向排序
        # 结构化输出
        if keep_format == "0":
            return_result = result
        elif keep_format == "2":
            for region in result:
                if (
                    region["type"].lower() == "table"
                    and len(region["res"]) > 0
                    and "html" in region["res"]
                ):
                    return_result.extend(
                        get_pos_contents(
                            region["res"],
                            from_source="structure-table",
                            draw_img=oneimg_c,
                        )
                    )
                    html_content.append(region["res"]["html"])
                else:
                    return_result.extend(
                        get_pos_contents(
                            region["res"], from_source="structure", draw_img=oneimg_c
                        )
                    )
        return oneimg_c, return_result, html_content

    # 把图片识别出内容,并保存到相应目录
    # 1、保存在每张图片对应的子目录下,
    # 仿照save_structure_res逻辑重写，理由是，这个方法是一个图片一个目录来存放结果报文、excel、figure,
    # 而重新定义的是一个图片相关信息通过文件名目录来体现信息来自同一图片，所有信息都放在同一目录info_dir下，就这点区别
    # 2、仿照convert_info_docx重写版面恢复成word逻辑，一则里面的逻辑基于save_structure_res，二则为了减少迭代识别res
    # 的迭代次数
    def infopic_info_save(self, oneimg, info_dir, info_filename_prefix, save_pdf=False):
        oneimg_c = oneimg.copy()
        # _, result, _ = self.ocr_structure(oneimg_c, keep_format="0")
        result = self.ocr_server.ocr_structure(
            oneimg_c, return_ocr_result_in_table=True
        )
        # 版面恢复
        # save_structure_res(result, self.output, '.')
        # for line in result:
        #     line.pop('img')
        #     print(line)
        h, w, _ = oneimg_c.shape
        res = sorted_layout_boxes(result, w)
        res_cp = deepcopy(res)
        # for word
        doc = Document()
        doc.styles["Normal"].font.name = "Times New Roman"
        doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        doc.styles["Normal"].font.size = shared.Pt(6.5)
        flag = 1
        with open(
            info_dir.joinpath(f"{info_filename_prefix}_res.txt"), "w", encoding="utf-8"
        ) as f:
            for region in res_cp:
                # for word begin
                if flag == 2 and region["layout"] == "single":
                    section = doc.add_section(WD_SECTION.CONTINUOUS)
                    section._sectPr.xpath("./w:cols")[0].set(qn("w:num"), "1")
                    flag = 1
                elif flag == 1 and region["layout"] == "double":
                    section = doc.add_section(WD_SECTION.CONTINUOUS)
                    section._sectPr.xpath("./w:cols")[0].set(qn("w:num"), "2")
                    flag = 2
                # for word end
                roi_img = region.pop("img")
                pic_fig_region = ".".join([str(i) for i in region["bbox"]])
                if (
                    region["type"].lower() == "table"
                    and len(region["res"]) > 0
                    and "html" in region["res"]
                ):
                    # 表格图片区域--下面尽管变量没用但需要pop,否则因为存在ndarray,f.write(..这一句会报错
                    region["res"].pop("boxes")
                    # table_img = region['res'].pop('boxes')
                    # 表格转化成exce
                    excel_path = info_dir.joinpath(
                        f"{info_filename_prefix}_{pic_fig_region}.xlsx"
                    )
                    to_excel(region["res"]["html"], excel_path)
                    # for word
                    paragraph = doc.add_paragraph()
                    new_parser = HtmlToDocx()
                    new_parser.table_style = "TableGrid"
                    table = new_parser.handle_table(html=region["res"]["html"])
                    new_table = deepcopy(table)
                    new_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    paragraph.add_run().element.addnext(new_table._tbl)
                elif region["type"].lower() == "figure":
                    img_path = info_dir.joinpath(
                        f"{info_filename_prefix}_{pic_fig_region}.jpg"
                    )
                    cv2.imencode(".jpg", roi_img)[1].tofile(img_path)
                    # for word
                    paragraph_pic = doc.add_paragraph()
                    paragraph_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph_pic.add_run("")
                    if flag == 1:
                        run.add_picture(str(img_path), width=shared.Inches(5))
                    elif flag == 2:
                        run.add_picture(str(img_path), width=shared.Inches(2))
                elif region["type"].lower() == "title":
                    doc.add_heading(region["res"][0]["text"])
                else:
                    paragraph = doc.add_paragraph()
                    paragraph_format = paragraph.paragraph_format
                    for i, line in enumerate(region["res"]):
                        if i == 0:
                            paragraph_format.first_line_indent = shared.Inches(0.25)
                        text_run = paragraph.add_run(line["text"] + " ")
                        text_run.font.size = shared.Pt(10)
                # 最后把整个信息保存到文件中，ensure_ascii=False是为了避免像中文中这种“\u51fa\u751f\u5730”ASCII编码
                f.write(f"{json.dumps(region, ensure_ascii=False)}\n")
        # save to docx
        docx_path = info_dir.joinpath(f"{info_filename_prefix}_word.docx")
        doc.save(docx_path)
        logging.information(f"docx save to {docx_path}")
        # save to pdf
        if save_pdf:
            pdf_path = info_dir.joinpath(f"{info_filename_prefix}_pdf.pdf")
            convert(docx_path, pdf_path)
            logging.information(f"pdf save to {pdf_path}")
        return result, oneimg


"""
    其他脚本直接引用此实例化的对象可以实现单例模型，除非特别需要单独实例化，
    可以一直引用此对象，保证整个程序只有一个实例，节约系统资源,不用轻易实例化
    上面实例，如果后期需要多线程，再思考多实例问题
"""
# ocr_request = OCRRequest()

# ocr_request_process = OCRRequest()

# if __name__ == '__main__':
# ocr = PaddleOCR(use_angle_cls=True, lang='ch') # need to run only once to download and load model into memory
# img_path = os.path.join(dirName,'01.png')
# result = ocr.ocr(img_path, cls=True)
# for line in result:
#     print(line)


#### 2、定义client用于访问子进程中的服务Process
# 识别后返回格式统一为  img,list,前者为图片，上面带有识别文字的标记框，后者是  \n分割的str
class OCRRequestProcess:
    def __init__(self, client_end_conn=None):
        # 进程通信间的变量
        self.client_end_conn = client_end_conn

    # 发送图片并获取ocr结果
    @clock
    def ocr(self, cv2_img):
        cv2_img_c = cv2_img.copy()
        # ocr方法支持 传一个图片列表进行批量识别，所以返回的[result1,result2.....]
        # 所以注意单个图片时也是一个[result],目前应用都是一个图片，所以下面结果直接取returnResult[0]
        # 后面根据需要，再使用批量图片上传的功能
        self.client_end_conn.send((cv2_img_c, "ocr", None))
        result = self.client_end_conn.recv()  # 只要服务端不关闭，就一直阻塞等待返回的消息
        return_list = get_pos_contents(result, from_source="ocr", draw_img=cv2_img_c)
        # return_str = "\n".join([str(line[4]) for line in returnResult])
        return cv2_img_c, return_list

    # 发送图片并获取ocr结果
    # -------------------图片识别内容还没有解决
    @clock
    def ocr_structure(self, cv2_img):
        cv2_img_c = cv2_img.copy()
        # 注意：send()和recv()方法使用pickle模块对对象进行序列化
        self.client_end_conn.send((cv2_img_c, "structure", True))
        result = self.client_end_conn.recv()
        # print(result)
        # print("-------------------")
        # 先按y方向排序 在x方向排序 lambda x, y: f"{str(x) if x is not None else ''} {str(y) if y is not None else ''}", row.values.tolist()
        result_sorted = sorted(result, key=lambda x: (x["bbox"][1], x["bbox"][0]))
        # print(result_sorted)
        result_pos_content = []
        html_content = []
        # 为了打印看输出结构信息
        # print("-----------")
        # for oneitem in result:
        #     oneitem.pop("img")
        #     oneitem['res'].pop('cell_bbox')
        #     print(f'result--oneitem{oneitem}')
        result_list = []
        for region in result_sorted:
            # 清理下数据减少内存
            region.pop("img")
            if (
                region["type"].lower() == "table"
                and len(region["res"]) > 0
                and "html" in region["res"]
            ):
                # 清理下数据减少内存
                region["res"].pop("cell_bbox")
                # print("structure-table识别的原始信息：：----------------------------")
                # print(region)
                # **2假设一个图片文档，遇到表格table时，如何格式化并什么形式（比如pd），后期再讨论
                result_pos_content.extend(
                    get_pos_contents(
                        region["res"],
                        from_source="structure-table",
                        draw_img=cv2_img_c,
                        table_pos=region["bbox"],
                    )
                )
                html_content.append(region["res"]["html"])
            else:
                # print("structure-other识别的原始信息：：-------------------------")
                # print(region)
                result_pos_content.extend(
                    get_pos_contents(
                        region["res"], from_source="structure-other", draw_img=cv2_img_c
                    )
                )
        # **2剩余的非表格内容统一装到结果pd里
        # if result_pos_content:
        #     content_list = list(zip(*result_pos_content))[4]
        #     # print("-------------------------------------------------------")
        #     # print(content_list)
        #     # 去重，并恢复原来顺序
        #     result_list = sorted(set(content_list), key=content_list.index)
        #     print("++++++++++++++++++++++++++++++++++++++++++++++++++++++++++")
        #     print(result_list)
        #     # result_str = "\n".join(result_list)
        #     # print("222222222222222222222222222222222222222222222222222222222")
        #     # print(result_str)
        return cv2_img_c, result_pos_content, html_content


# 启动ocr识别服务
# 在线程中实例化 OCRRequestProcess，并存储为全局变量，全局可调用进行OCR，并在实例化时启动模型服务
def startup_ocrserver():
    # 进程通信间的变量
    client_end_conn, server_end_conn = Pipe()
    # 开一个守护进程来作为服务，运行在整个程序期间
    worker_ocr_server = Worker(
        "OCRServerProcess",
        classMethod="start",
        classMethodArgs={},
        module="model.OCRServer",
        serverend_conn=server_end_conn,
        name="anne",
        daemon=True,
    )
    QThreadPool.globalInstance().start(worker_ocr_server)

    worker_ocr_req = Worker(
        "OCRRequestProcess", module="route.OCRRequest", client_end_conn=client_end_conn
    )
    worker_ocr_req.communication.result.connect(partial(globalvar.set_var, "ocrserver"))
    QThreadPool.globalInstance().start(worker_ocr_req)


# 调用ocr服务
# 在线程中实例化 OCRRequestProcess，并存储为全局变量，全局可调用进行OCR，并在实例化时启动模型服务
def get_ocr_result(
    img,
    call_back_fun,
    adjust_para=None,
    ocr_method="ocr",
):
    worker = Worker(ocr, img, adjust_para, ocr_method)
    worker.communication.result.connect(call_back_fun)
    QThreadPool.globalInstance().start(worker)


def ocr(img: np.ndarray | Image.Image | str, adjust_para=None, ocr_method="ocr"):
    if isinstance(img, str):
        img = Image.open(str)
    if adjust_para is not None:
        img = scale_image(img, adjust_para)
    if isinstance(img, Image.Image):
        # 把图片转化为numpy的格式
        img = cv2.cvtColor(np.asarray(img), cv2.COLOR_RGB2BGR)
        # img = cv2.imdecode(
        #     np.asarray(bytearray(img), dtype=np.uint8), cv2.IMREAD_COLOR
        # )
    ocr_request = globalvar.get_var("ocrserver")
    while not ocr_request:
        time.sleep(1)
        # QThread.msleep(1)
        ocr_request = globalvar.get_var("ocrserver")
    final_cv2_draw = None
    final_info = []
    final_html_content = []
    h, w, d = img.shape
    # 把长图片拆成小图片，设置成图片高度超过1920就拆图片再识别，然后再合并图片和识别内容
    # 谨慎使用，因为如果仅从高度上拆分，如果宽度很大，模型会自动缩放宽度，并等比例缩放高度，导致高度会非常小，这样就无法ocr
    # 下面拆分逻辑为，如果图片高度超过1920应该分拆图片，但分拆前提是
    # 如果宽度很多，远远超过1920，模型会压缩到1920，这时等比例高度收缩不能太小（设置成680），否则会导致形成窄窄一条，无法识别
    # h / (w / 1920)表示w收缩导致h最终收缩到高度
    # 1080为模型最长边最大长度是参数det_limit_side_len的设置
    h_min_unit_size = (w / 1080) * 680  # 高度分拆的单位最小高度
    split_num = h // h_min_unit_size  # 高度能分的块数
    if w / 1080 > 1 and split_num > 1:
        # 以1920为单位 把长图片分成小图片
        print(f"get_ocr_result，图片高度超过1920自动拆成小图片再识别")
        imgs_list = image_split.split(img, split_length=int(split_num), axis=1)
        for img_cv2_seg in imgs_list:
            # 一小段一小段识别
            if ocr_method == "ocr":
                img_cv2_seg_draw, out_info = ocr_request.ocr(img_cv2_seg)
            else:
                img_cv2_seg_draw, out_info, html_content = ocr_request.ocr_structure(
                    img_cv2_seg
                )
                final_html_content.extend(html_content)
            # cv2.imwrite(f'{Path(__file__).parent}/tmp/seg{i}.png',img_cv2_seg_draw)
            # 画识别的区域的图片再拼接起来
            final_cv2_draw = (
                img_cv2_seg_draw
                if final_cv2_draw is None
                else vstack((final_cv2_draw, img_cv2_seg_draw))
            )
            # 识别内容存起来
            final_info.extend(out_info)
    else:
        if ocr_method == "ocr":
            final_cv2_draw, final_info = ocr_request.ocr(img)
        else:
            final_cv2_draw, final_info, html_content = ocr_request.ocr_structure(img)
    # cv2.imwrite(f'{Path(__file__).parent}/tmp/ocr_draw_result.png', final_cv2_draw)
    return img, final_cv2_draw, final_info, final_html_content


# 3、通过IP访问独立服务器
class OCRRequestNet:
    def upload_ocr_info_ip(self, qPixmap):
        shot_bytes = get_shot_bytes(qPixmap)
        filename = "shot" + str(time.time()).split(".")[0] + ".jpg"
        files = {"file": (filename, shot_bytes, "image/jpeg")}
        cfg = ConfigParser()
        cfg.read("config.ini")
        headers = {"Cookie": cfg.get("picbed", "cookie")}
        try:
            res = requests.post(
                cfg.get("picbed", "api"),
                data={"compress": 960},
                files=files,
                headers=headers,
            )
            status_code = res.json()["code"]
            if status_code == 200:
                return res.json()["data"]["raw_out"]
            else:
                logging.information("服务非正常返回", f"not 200, code: {res.status_code}")
        except Exception as e:
            print(e.args)
