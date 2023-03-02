from PySide6.QtCore import Qt, QEvent
from PySide6.QtWidgets import (
    QApplication,
    QDialog,
    QFileDialog,
    QTableView,
    QHeaderView,
)
from mytools.video.model.inference.OCRVideo import OCRVideo
from myutils.image_convert import cv2pixmap
from .OCRResult import Ui_OCRResult
import pandas as pd
import numpy as np
import cv2
from myutils.office_tools import to_excel_auto_column_weight
from ui.component.pandas_model import PandasModel
from functools import reduce
from docx import Document


class TotalMessage(QDialog, Ui_OCRResult):
    def __init__(self, content):
        super().__init__()
        self.setupUi(self)
        # self.adjustSize()
        self.setWindowFlags(Qt.WindowStaysOnTopHint)
        # self.setWindowFlag(Qt.FramelessWindowHint)  # 没有窗口栏
        self.setAttribute(Qt.WA_TranslucentBackground)  # 设置背景透明
        self.select_btn_pic.clicked.connect(self.ocr_pdf)
        self.select_btn_pdf.clicked.connect(self.ocr_pdf)
        self.start_btn_scan.clicked.connect(self.ocr_pdf)
        self.start_btn_video.clicked.connect(self.start_ocr_video)
        self.start_btn_video_screenshot.clicked.connect(self.ocr_video_screenshot)
        self.start_btn_video_monitor.clicked.connect(self.ocr_video_monitor)
        self.start_btn_video_stop.clicked.connect(self.ocr_video_close)
        self.table_btn.clicked.connect(self.display_data_table)
        self.export_data_excel.clicked.connect(self.save_data_excel)
        self.export_data_word.clicked.connect(self.save_data_word)
        self.gridEdit.setAlignment(Qt.AlignCenter)
        self.setWindowFlag(Qt.Tool)  # 不然exec_执行退出后整个程序退出
        self.init_info(content)
        self.scrollArea.installEventFilter(self)
        self.last_time_move = 0

    def init_info(self, content):
        if type(content) is np.ndarray:
            show = cv2.resize(content, (320, 260))
            self.cameraLabel.setPixmap(cv2pixmap(show))
            # 以下是使用QTextEdit方式插入图片，这个会不断往下追加，先用QLabel吧
            # document = self.videoEdit.document()
            # cursor = QTextCursor(document)
            # p1 = cursor.position() # returns int
            # cursor.insertImage(showImage)
        if type(content) is tuple:
            # 识别的图，上面带识别的区域划线
            # show = cv2.resize(content[0],(480,300))
            width = int(self.cameraLabel.width() * 1.1)
            height = int(self.cameraLabel.height() * 1.1)
            m_pixmap = cv2pixmap(content[0])
            # self.scrollAreaWidgetContents.resize(width, height)
            # m_pixmap = m_pixmap.scaled(width, height, Qt.KeepAspectRatio, Qt.SmoothTransformation)
            self.cameraLabel.setPixmap(m_pixmap)
            # 识别的内容列表
            self.content_pd = pd.DataFrame(content[1])
            # to_s = self.content_pd.to_string(index=False,header=False, justify='left')
            df_series = self.content_pd.apply(
                lambda row: reduce(
                    lambda x, y: f"{str(x) if x is not None else ''} {str(y) if y is not None else ''}",
                    row.values.tolist(),
                ),
                axis=1,
            )
            to_s = "\n".join(df_series.values)
            self.gridEdit.setText(to_s)
            # self.gridEdit.setAlignment(Qt.AlignLeft)
            # 如果原生带格式识别，就展示出来
            self.htmlEdit.clear()
            for html_script in content[2]:
                self.htmlEdit.append(html_script)
        if type(content) is str:
            self.gridEdit.setText(content)
        # self.cameraLabel.resize(480,260)

    def display_data_table(self):
        self.view = QTableView()
        self.view.resize(1200, 800)
        self.view.resizeColumnsToContents()
        self.view.horizontalHeader().setStretchLastSection(True)
        self.view.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)
        self.view.setAlternatingRowColors(True)
        self.view.setSelectionBehavior(QTableView.SelectRows)
        # self.view.resizeRowsToContents()
        # QTableWidget.resizeRowsToContents(self.view)
        model = PandasModel(self.content_pd)
        self.view.setModel(model)
        self.view.show()
        # self.hide()

    def save_data_excel(self):
        file, _ = QFileDialog.getSaveFileName(
            self, "保存到", "example.xlsx", "excel files(*.xlsx *xls)"
        )
        if file:
            # df = pd.read_csv(self.plainTextEdit.getPaintContext(), sep='\t')
            to_excel_auto_column_weight({"sheet01": self.content_pd}, file_name=file)
            # .to_excel(file)
        # self.close()
        # 选择文件进行pdf内容识别

    def save_data_word(self):
        file, _ = QFileDialog.getSaveFileName(
            self, "保存到", "example.docx", "word files(*.docx *doc)"
        )
        if file:
            # 创建 Document 对象，等价于在电脑上打开一个 Word 文档
            document = Document()
            # 在 Word 文档中添加一个标题
            document.add_heading("这是一个标题", level=0)
            # 文档添加段落
            p = document.add_paragraph("这是白给的段落")
            # 添加带样式的文字
            # 添加段落，文本可以包含制表符（\t）、换行符（\n）或回车符（\r）等
            # add_run() 在段落后面追加文本
            p.add_run("\n我倾斜了").italic = True  # 添加一个倾斜文字
            p.add_run(f"\n{self.gridEdit.toPlainText()}").bold = True  # 添加一个加粗文字
            # 保存文档
            document.save(file)
        # self.close()
        # 选择文件进行pdf内容识别

    def ocr_pdf(self):
        file, _ = QFileDialog.getOpenFileName(
            self, "选择文件", "screenshot.jpg", "pdf files(*.pdf)"
        )
        if file:
            shot_img = self.cameraLabel.pixmap().toImage()
            shot_img.save(file)
        self.close()

    # 视频扫描
    def start_ocr_video(self):
        if not hasattr(self, "video_ocr"):
            self.video_ocr = OCRVideo()
            self.video_ocr.signal.connect(self.ocr_video_result)
            self.video_ocr.signal_img.connect(self.ocr_video_result)
            # # 准备就绪，启动线程开始扫描 此时 run中
            self.video_ocr.start()
        if not self.video_ocr.isRunning():
            self.video_ocr.start()

    def ocr_video_screenshot(self):
        self.video_ocr.pause()

    def ocr_video_monitor(self):
        self.video_ocr.resume()

    def ocr_video_close(self):
        self.video_ocr.stop()

    def ocr_video_result(self, result):
        # print(result)
        self.init_info(result)

    def save_local(self):
        file, _ = QFileDialog.getSaveFileName(
            self, "保存到", "screenshot.jpg", "Image files(*.jpg *.gif *.png)"
        )
        if file:
            shot_img = self.cameraLabel.pixmap().toImage()
            shot_img.save(file)
        self.close()

    def save_to_clipboard(self):
        clipboard = QApplication.clipboard()
        clipboard.setPixmap(self.cameraLabel.pixmap())
        self.close()

    def closeEvent(self, event):
        if hasattr(self, "video_ocr") and self.video_ocr:
            self.video_ocr.stop()

    def open_image(self):
        """
        select image file and open it
        :return:
        """
        img_name, _ = QFileDialog.getOpenFileName(
            self, "Open Image File", "*.jpg;;*.png;;*.jpeg"
        )
        self.box.set_image(img_name)

    # 成员函数copyPic
    # 作用：复制图片
    def copyPic(self):
        # 打开源图片
        # self.imageName表示图片路径，为字符串。（是在Ui_MainWindow类的构造函数__init__(self)中定义的成员变量。）
        f_src = open(self.cameraLabel, "rb")
        # 读取图片内容并存储到content变量
        content = f_src.read()

        # 以二进制格式打开复制后的图片（只写）
        # wb一般用于非文本文件如图片等。
        # 如果该文件已存在则打开文件，并从开头开始编辑，即原有内容会被删除。
        # 如果该文件不存在，创建新文件。
        f_copy = open("./1.jpg", "wb")

        # 源图片的内容以二进制形式写入新图片
        f_copy.write(content)
        # 关闭文件（原则：先打开的后关闭）
        f_copy.close()
        f_src.close()

    def eventFilter(self, source, event):
        # 双击拷贝图片
        print(event.type())
        if event.type() == QEvent.MouseButtonDblClick:
            clipboard = QApplication.clipboard()
            clipboard.setPixmap(self.cameraLabel.pixmap())
        elif (
            event.type() == QEvent.MouseButtonPress and event.button() == Qt.LeftButton
        ):
            """
            mouse press events for the widget
            :param e: QMouseEvent
            :return:
            """
            self.left_click = True
            self.start_pos = event.pos()
        elif event.type() == QEvent.MouseMove and event.button() == Qt.LeftButton:
            """
            mouse move events for the widget
            :param event: QMouseEvent
            :return:
            """
            self.end_pos = event.pos() - self.start_pos
            self.point = self.point + self.end_pos
            self.start_pos = event.pos()
            self.repaint()
        elif (
            event.type() == QEvent.MouseButtonRelease
            and event.button() == Qt.LeftButton
        ):
            """
            mouse release events for the widget
            :param e: QMouseEvent
            :return:
            """
            self.left_click = False
